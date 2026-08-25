"""素材文本抽取：PDF(文本层) / DOCX / MD / TXT → 文本块列表。

每个文本块 = {index, label(页码/来源), text, chars}；扫描件 PDF 显式报错并给出提示。
"""

from pathlib import Path
from typing import Any


class ExtractError(Exception):
    """素材解析失败（含可读提示）。"""


def extract_text(path: str | Path) -> list[dict[str, Any]]:
    p = Path(path)
    if not p.exists():
        raise ExtractError(f"文件不存在: {p}")
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix == ".docx":
        return _extract_docx(p)
    if suffix in (".md", ".markdown", ".txt"):
        return _extract_textfile(p)
    raise ExtractError(f"不支持的文件类型: {suffix}（支持 PDF/DOCX/MD/TXT）")


def _extract_pdf(p: Path) -> list[dict[str, Any]]:
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ExtractError("未安装 PyMuPDF（pip install pymupdf）") from e
    blocks: list[dict[str, Any]] = []
    total = 0
    with fitz.open(str(p)) as doc:
        for i, page in enumerate(doc):
            text = (page.get_text() or "").strip()
            if not text:
                continue
            total += len(text)
            blocks.append({"index": len(blocks), "label": f"P{i + 1}", "text": text,
                           "chars": len(text)})
    if total < 200:
        raise ExtractError(
            "该 PDF 疑似扫描件（无文本层）。请使用带文本层的 PDF、或先用 WPS/OCR 转成文字文档。"
        )
    return blocks


def _extract_docx(p: Path) -> list[dict[str, Any]]:
    try:
        import docx
    except ImportError as e:
        raise ExtractError("未安装 python-docx（pip install python-docx）") from e
    d = docx.Document(str(p))
    parts: list[str] = [para.text for para in d.paragraphs if para.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if len(text) < 50:
        raise ExtractError("DOCX 内容过少或为空")
    return [{"index": 0, "label": "DOCX", "text": text, "chars": len(text)}]


def _extract_textfile(p: Path) -> list[dict[str, Any]]:
    raw = p.read_bytes()
    text = None
    for enc in ("utf-8", "gbk"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = raw.decode("utf-8", errors="replace")
    if len(text.strip()) < 20:
        raise ExtractError("文本文件内容过少或为空")
    return [{"index": 0, "label": p.suffix.upper(), "text": text, "chars": len(text)}]
